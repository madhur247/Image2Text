from flask import *
import cv2
import pytesseract
import shutil, os
import base64
from flask_mail import Mail, Message
import random
import secrets
from flask_mysqldb import MySQL
from flask_session import Session
from datetime import datetime
from docx import Document
import re
from docx.shared import Pt
from fpdf import FPDF
import bcrypt

app = Flask(__name__)
app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = ''
app.config['MYSQL_DB'] = 'newspaper_recog_userdb'

mysql = MySQL(app)
app.secret_key = secrets.token_bytes(16)
app.config['SESSION_PERMANENT'] = False
app.config['SESSION_TYPE'] = "filesystem"
Session(app)
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USERNAME'] = 'image2text123@gmail.com'
app.config['MAIL_PASSWORD'] = '*************'
app.config['MAIL_USE-TLS'] = False
app.config['MAIL_USE_SSL'] = True
mail = Mail(app)


def createdocx(output):
    document = Document()
    myfile = re.sub(r'[^\x00-\x7F]+|\x0c', ' ', output) # remove all non-XML-compatible characters
    para = document.add_paragraph().add_run(myfile)
    para.font.size = Pt(12)
    document.save('.\static\output\output.docx')

def createpdf(output):

    pdf = FPDF()

    pdf.add_page()


    pdf.set_font("Times", size=12)


    f = open(".\static\output\output.txt", "r")

    for x in f:
        x = re.sub(r'[^\x00-\x7F]+|\x0c', ' ', x)
        pdf.cell(200, 10, txt=x, ln=1, align='L')

    pdf.output(".\static\output\output.pdf")


def predict_txt(img_path):

    pytesseract.pytesseract.tesseract_cmd = 'C:\Program Files\Tesseract-OCR\\tesseract.exe'

    img1 = cv2.imread(img_path)
    img = cv2.cvtColor(img1, cv2.COLOR_BGR2GRAY)
    # img = cv2.adaptiveThreshold(img1, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 101, 5)

    output = pytesseract.image_to_string(img)

    file = open(".\static\output\output.txt", "w")
    file.writelines(output)
    file.close()
    
    createdocx(output)
    
    createpdf(output)
    
    return output

@app.route('/')
def index():
    session.clear()
    session['login_status'] = False
    return render_template("index.html", status=session.get('login_status'))

@app.route('/login_index')
def login_index():
    if session.get('login_status'):
        return render_template("index.html", status=session.get('login_status'),user=session.get('name'), plan = session.get('pack'))
    else:
        return render_template("index.html", status=session.get('login_status'))

@app.route('/login/<dest>')
def login(dest):
    session['login_dest'] = dest
    return render_template("login.html")

@app.route('/verify', methods=['GET','POST'])
def verify():
    if request.method == "POST":
        if request.form['name']!="" and request.form['email']!="" and request.form['password']!="" and request.form['con_password']!="":
            if request.form['password']==request.form['con_password']:
                ver_email = request.form['email']
                session['email'] = ver_email
                
                reg = None
                try:
                    cur = mysql.connection.cursor()
                    cur.execute("SELECT * FROM users WHERE email = '"+session.get('email')+"'")
                    reg = cur.fetchone()
                    print(reg,session.get('email'))
                    mysql.connection.commit()
                    cur.close()
                except:
                    print(reg,session.get('email'))
                if reg is None:
                    otp_gen = random.randint(100000,999999)
                    session['otp'] = otp_gen
                    print("verify")
                    print(session.get('otp'))
                    msg="Your One Time Password for Image2Text is: "+str(session.get('otp'))
                    message = Message("OTP Verification",sender="email",recipients=[session.get('email')])
        
                    message.body = msg
        
                    mail.send(message)
        
                    success = "Message Sent"
        
                    session['name'] = request.form.get('name')
                    session['hashed'] = bcrypt.hashpw(request.form['password'].encode('utf-8'), bcrypt.gensalt())
                    print("hash",session.get('hashed'))
                    return render_template("otp.html",success=success)
                else:
                    flash('This Email is already registered')
                    return render_template("login.html")
            else:
                flash('Password and Confirm Password Do not Match!')
                return render_template("login.html")
        else:
            flash('Please enter all the details properly!')
            return render_template("login.html")

@app.route('/authenticate', methods=['GET','POST'])
def authenticate():
    if request.method == "POST":
        otp_get = request.form.getlist('otp')
        otp_enter = ""
        for ele in otp_get:
            otp_enter+=str(ele)
        if len(otp_enter)<1:
            otp_enter = -1
        if int(otp_enter)==session.get('otp'):
            return render_template("pack.html")
        else:
            flash("Incorrect OTP\nTry Again!")
            return render_template("otp.html")

@app.route('/payment/<plan>', methods=['GET','POST'])
def payment(plan):
    if request.method == "POST":
        if plan=="basic":
            session['pack'] = "basic"
            return render_template("payment.html",amount="5")
        elif plan=="pro":
            session['pack'] = "pro"
            return render_template("payment.html",amount="25")
        elif plan=="premium":
            session['pack'] = "premium"
            return render_template("payment.html",amount="40")
        elif plan=="free":
            session['pack'] = "free"
            reg = None
            try:
                cur = mysql.connection.cursor()
                cur.execute("SELECT * FROM users WHERE email = '"+session.get('email')+"'")
                reg = cur.fetchone()
                print(reg)
                mysql.connection.commit()
                cur.close()
            except:
                print(reg,session.get('email'))
            if reg is None:
                cur = mysql.connection.cursor()
                print(str(session.get('hashed')))
                cur.execute("INSERT INTO users VALUES('"+session.get('name')+"','"+session.get('email')+"','"+str(datetime.now())+"','"+str(session.get('hashed'))[2:-1]+"','"+session.get('pack')+"')")
                mysql.connection.commit()
                cur.close()
            else:
                cur = mysql.connection.cursor()
                cur.execute("UPDATE users SET Start_date='"+str(datetime.now())+"', Package='"+session.get('pack')+"' WHERE Email='"+session.get('email')+"'")
                mysql.connection.commit()
                cur.close()
            session['login_status'] = True
            if session.get('login_dest')=="index" or session.get('login_dest')=="pack":
                return render_template("index.html",status=session.get('login_status'),user=session.get('name'), plan = session.get('pack'))
            elif session.get('login_dest')=="upload":
                return render_template("upload.html",user=session.get('name'))
     
@app.route('/contact',methods=['POST','GET'])
def contact():
    if request.method=="POST":
        name = request.form['feed_name']
        email = request.form['feed_email']
        comment = request.form['comments']
        cur = mysql.connection.cursor()
        cur.execute("INSERT INTO feedback VALUES('"+name+"','"+email+"','"+comment+"')")
        mysql.connection.commit()
        cur.close()
        return render_template("feedback.html")
    
@app.route('/signup_index',methods=['GET','POST'])
def signup_index():
    if request.method == "POST":
        reg = None
        try:
            cur = mysql.connection.cursor()
            cur.execute("SELECT * FROM users WHERE email = '"+session.get('email')+"'")
            reg = cur.fetchone()
            print(reg)
            mysql.connection.commit()
            cur.close()
        except:
            print(reg,session.get('email'))
        if reg is None:
            cur = mysql.connection.cursor()
            cur.execute("INSERT INTO users VALUES('"+session.get('name')+"','"+session.get('email')+"','"+str(datetime.now())+"','"+str(session.get('hashed'))[2:-1]+"','"+session.get('pack')+"')")
            mysql.connection.commit()
            cur.close()
        else:
            cur = mysql.connection.cursor()
            cur.execute("UPDATE users SET Start_date='"+str(datetime.now())+"', Package='"+session.get('pack')+"' WHERE Email='"+session.get('email')+"'")
            mysql.connection.commit()
            cur.close()
        session['login_status'] = True
        if session.get('login_dest')=="index" or session.get('login_dest')=="pack":
            return render_template("index.html",status=session.get('login_status'),user=session.get('name'), plan = session.get('pack'))
        elif session.get('login_dest')=="upload":
            return render_template("upload.html",user=session.get('name'))
            
@app.route('/resend_verify', methods=['GET','POST'])
def resend_verify():
    if request.method == "POST":
        #email = request.form['email']
        otp_gen = random.randint(100000,999999)
        session['otp'] = otp_gen
        print("resend")
        print(session.get('otp'))
        msg="Your One Time Password for Image2Text is: "+str(session.get('otp'))
        message = Message("OTP Verification",sender='image2text123@gmail.com',recipients=[session.get('email')])

        message.body = msg

        mail.send(message)

        success = "Message Sent"
        
        flash(u"New OTP is sent to {0}".format(session.get('email')))
        return render_template("otp.html")

# @app.route('/selection',methods=['GET','POST'])
# def selection():
#     if request.method=='POST':
#         chkbox = request.form.getlist('chkbox')
#         session['profile'] = chkbox[0]
#         cur = mysql.connection.cursor()
#         cur.execute("INSERT INTO users VALUES('"+session.get('name')+"','"+session.get('email')+"','"+str(datetime.now())+"','"+session.get('password')+"','"+session.get('profile')+"')")
#         mysql.connection.commit()
#         cur.close()
#         if session.get('profile')=="Free":
#             return render_template("index.html",user_name=session.get('name'),prof=session.get('profile'))
#         elif session.get('profile')=="Paid":
#             return render_template("index.html",user_name=session.get('name'),prof=session.get('profile'))
        
@app.route('/userlogin',methods=['GET','POST'])
def userlogin():
    if request.method=='POST':
        if request.form['loginemail'] != "" and request.form['loginpass'] != "":
            email = request.form['loginemail']
            password = request.form['loginpass']
            reg = None
            try:
                cur = mysql.connection.cursor()
                cur.execute("SELECT * FROM users WHERE email = '"+email+"'")
                reg = cur.fetchone()
                session['hashed'] = reg[3].encode("utf-8")
                session['email'] = reg[1]
                session['name'] = reg[0]
                session['pack'] = reg[4]
                session['Start_date'] = reg[2]
                mysql.connection.commit()
                cur.close()
            except:
                pass
            if reg is None:
                flash('This Email is not registered')
                return render_template("login.html")
            else:
                if bcrypt.checkpw(password.encode('utf-8'), session.get('hashed')):
                    pack = session.get('pack')
                    print(pack)
                    today_date = str(datetime.now())[:10]
                    date_format = "%Y-%m-%d"
                    a = datetime.strptime(today_date, date_format)
                    b = datetime.strptime(session.get('Start_date'), date_format)
                    delta = a-b
                    
                    if pack=="free":
                        if delta.days>15:
                            flash('Your Free Trial has Ended')
                            return render_template("pack2.html")
                        else:
                            session['login_status'] = True
                            if session.get('login_dest')=="index":
                                return render_template("index.html",status=session.get('login_status'),user=session.get('name'), plan = session.get('pack'))
                            elif session.get('login_dest')=="upload":
                                return render_template("upload.html",user=session.get('name'))
                            elif session.get('login_dest')=="pack":
                                flash("Your Free Trial is active")
                                return render_template("pack2.html")
    
                    elif pack=="basic":
                        if delta.days>30:
                            flash('Your Plan has Ended')
                            return render_template("pack2.html")
                        else:
                            session['login_status'] = True
                            if session.get('login_dest')=="index":
                                return render_template("index.html",status=session.get('login_status'),user=session.get('name'), plan = session.get('pack'))
                            elif session.get('login_dest')=="upload":
                                return render_template("upload.html",user=session.get('name'))
                            elif session.get('login_dest')=="pack":
                                flash("Your Basic Plan is active")
                                return render_template("pack2.html")
                    
                    elif pack=="pro":
                        if delta.days>180:
                            flash('Your Plan has Ended')
                            return render_template("pack2.html")
                        else:
                            session['login_status'] = True
                            if session.get('login_dest')=="index":
                                return render_template("index.html",status=session.get('login_status'),user=session.get('name'), plan = session.get('pack'))
                            elif session.get('login_dest')=="upload":
                                return render_template("upload.html",user=session.get('name'))
                            elif session.get('login_dest')=="pack":
                                flash("Your Pro Plan is active")
                                return render_template("pack2.html")
                    
                    elif pack=="premium":
                        if delta.days>365:
                            flash('Your Plan has Ended')
                            return render_template("pack2.html")
                        else:
                            session['login_status'] = True
                            if session.get('login_dest')=="index":
                                return render_template("index.html",status=session.get('login_status'),user=session.get('name'), plan = session.get('pack'))
                            elif session.get('login_dest')=="upload":
                                return render_template("upload.html",user=session.get('name'))
                            elif session.get('login_dest')=="pack":
                                flash("Your Premium Plan is active")
                                return render_template("pack2.html")
                else:
                    flash('Incorrect Password!')
                    return render_template("login.html")
        else:
            flash('Please enter all the details properly!')
            return render_template("login.html")
            
@app.route('/upload',methods=['GET','POST'])
def upload():
    if request.method=="POST":
        return render_template("upload.html",user=session.get('name'))

@app.route('/convert', methods=['GET', 'POST'])
def convert():
    if request.method == 'POST':

        file = request.files['file']
        file1 = request.form['base64data']

        if file:
            # print("path" + os.path.join(os.path.curdir, secure_filename(file.filename)))
            shutil.rmtree('.\static\images1')
            os.makedirs('.\static\images1')

            img_path = os.path.join('.\static\images1', "input_image.png")

            file.save(img_path)

            output_text = predict_txt(img_path)
            return render_template("download.html", img=str(img_path), op_text=str(output_text), user=session.get('name'), pack=session.get('pack'))

        elif file1:
            file1 += '=' * (-len(file1) % 4)
            with open(".\static\images1\input_image.png", "wb") as fh:
                fh.write(base64.decodebytes(file1.split(',')[1].encode()))
            img_path = ".\static\images1\input_image.png"
            
            output_text = predict_txt(img_path)

            return render_template("download.html", img=str(img_path), op_text=str(output_text), user=session.get('name'), pack=session.get('pack'))

        else:
            return "No Image File"

if __name__ == '__main__':
    app.run()
